from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
SUBMISSION_DIR = ROOT / "submission"
WRITEUP_PATH = SUBMISSION_DIR / "Agentic_Healthcare_Assistant_Writeup.docx"
SOURCE_ZIP_PATH = SUBMISSION_DIR / "Agentic_Healthcare_Assistant_Source_Code.zip"


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(7)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def create_writeup() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    document.add_heading("Agentic Healthcare Assistant for Medical Task Automation", 0)
    add_paragraph(document, "Project Write-up")

    document.add_heading("1. Project Overview", level=1)
    add_paragraph(
        document,
        "This project implements an agentic healthcare assistant that automates "
        "three routine healthcare tasks: patient-history lookup, doctor appointment "
        "scheduling, and retrieval-grounded medical-information responses. A Streamlit "
        "dashboard provides a simple interface for users to access all features."
    )

    document.add_heading("2. Objectives", level=1)
    add_bullets(document, [
        "Retrieve patient history using a patient ID or patient name.",
        "Check doctor availability and book valid appointment slots.",
        "Provide medical information only from approved reference documents.",
        "Route requests to the correct specialised agent.",
        "Protect patients by using a medical disclaimer and avoiding diagnosis or prescriptions.",
    ])

    document.add_heading("3. System Architecture", level=1)
    add_bullets(document, [
        "Planner Agent: classifies a request as patient-history, appointment, or medical-information work.",
        "History Agent: searches patient_db.csv and returns relevant patient details.",
        "Appointment Agent: validates patient and doctor IDs, checks schedules, prevents booked-slot conflicts, and updates appointments.csv.",
        "Medical Agent: retrieves approved document chunks before responding and includes a safety disclaimer.",
        "RAG Layer: reads TXT and DOCX documents, creates sentence-transformer embeddings, and stores them in a FAISS index.",
        "Streamlit UI: presents separate tabs for patient history, appointments, medical information, and request routing.",
    ])

    document.add_heading("4. Technology Stack", level=1)
    add_bullets(document, [
        "Python 3.11", "Streamlit", "Pandas", "FAISS", "Sentence Transformers", "python-docx", "python-dotenv"
    ])

    document.add_heading("5. Data and Workflow", level=1)
    add_paragraph(
        document,
        "The project uses CSV files for patient records, doctor schedules, and appointments. "
        "For appointment booking, the system checks the doctor's available working days and "
        "time range, generates 30-minute slots, removes already-booked slots, and writes a "
        "confirmed appointment back to appointments.csv."
    )
    add_paragraph(
        document,
        "For medical information, the vector-store pipeline reads approved TXT and DOCX "
        "documents, splits their text into chunks, creates embeddings with all-MiniLM-L6-v2, "
        "and saves the searchable index and metadata locally. The retriever returns only "
        "relevant approved chunks to the medical agent."
    )

    document.add_heading("6. Functional Testing", level=1)
    add_bullets(document, [
        "Patient-history test: search for P002 or Lakshmi Devi to retrieve medical history, medication, allergies, last visit, and assigned doctor.",
        "Appointment test: check slots for D002 on an available date, select an available time, and confirm that a new appointment is written to appointments.csv.",
        "RAG test: ask 'Can this assistant diagnose a disease?' and confirm that the answer is grounded in the policy documents and includes a safety note.",
        "Planner test: route requests such as 'Show allergies for P002' and 'Book an appointment with Dr. Arjun Rao'.",
    ])

    document.add_heading("7. Safety and Privacy", level=1)
    add_paragraph(
        document,
        "The assistant is designed for information and task automation only. It does not "
        "diagnose diseases or prescribe medication. Medical responses include a disclaimer "
        "and advise users to seek qualified clinical support for personal or urgent concerns. "
        "The user interface exposes only the requested patient record rather than the full database."
    )

    document.add_heading("8. Limitations and Future Scope", level=1)
    add_bullets(document, [
        "The current medical knowledge base contains policy and placeholder documents for testing; production use requires vetted clinical references.",
        "Authentication, role-based access, audit logs, encryption, and a production database should be added before real-world deployment.",
        "Future improvements can add appointment rescheduling/cancellation, notifications, multilingual support, and clinical-system integration.",
    ])

    document.add_heading("9. How to Run", level=1)
    add_paragraph(document, "1. Create/activate the virtual environment and install requirements: pip install -r requirements.txt")
    add_paragraph(document, "2. Build the index: python rag/vector_store.py")
    add_paragraph(document, "3. Start the dashboard: python -m streamlit run ui/streamlit_app.py")

    document.add_heading("10. Conclusion", level=1)
    add_paragraph(
        document,
        "The Agentic Healthcare Assistant demonstrates how specialised agents, structured "
        "healthcare data, and retrieval-augmented generation can automate common medical "
        "administrative tasks while maintaining safety boundaries for medical information."
    )

    document.save(WRITEUP_PATH)


def create_source_zip() -> None:
    include_files = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        relative = path.relative_to(ROOT)
        parts = set(relative.parts)
        if parts & {".git", ".venv", ".vscode", "__pycache__", "submission"}:
            continue
        if path.name == ".env" or path.suffix == ".pyc":
            continue
        if "faiss_index" in parts:
            continue
        include_files.append(path)

    with ZipFile(SOURCE_ZIP_PATH, "w", ZIP_DEFLATED) as archive:
        for path in include_files:
            archive.write(path, path.relative_to(ROOT))


if __name__ == "__main__":
    SUBMISSION_DIR.mkdir(exist_ok=True)
    create_writeup()
    create_source_zip()
    print(f"Created: {WRITEUP_PATH}")
    print(f"Created: {SOURCE_ZIP_PATH}")
